"""Tests for the export pipeline — Markdown, DOCX, and PDF exporters."""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from resumeforge.data.schema import (
    Bullet,
    Certification,
    Certifications,
    Education,
    EducationEntry,
    Experience,
    JobDescription,
    Meta,
    Position,
    Profile,
    Project,
    Projects,
    ResumeContext,
    SkillCategory,
    Skills,
)
from resumeforge.export import get_exporter
from resumeforge.export.base import ExportError
from resumeforge.export.docx_export import DocxExporter
from resumeforge.export.markdown import MarkdownExporter
from resumeforge.export.pdf import PdfExporter

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

TEMPLATE_DIR = (
    Path(__file__).parent.parent
    / "resumeforge"
    / "templates"
    / "classic"
)


def _make_context(locale: str = "en") -> ResumeContext:
    """Build a minimal but complete ResumeContext for testing."""
    return ResumeContext(
        profile=Profile(
            name="Jane Doe",
            title="Security Engineer",
            email="jane@example.com",
            phone="+1-555-0100",
            linkedin="linkedin.com/in/janedoe",
            location="Tel Aviv, Israel",
            summary="Experienced security engineer with a passion for cloud security.",
        ),
        experience=Experience(
            positions=[
                Position(
                    company="Acme Corp",
                    title="Senior Security Engineer",
                    start_date="2022-01",
                    end_date=None,
                    is_current=True,
                    location="Remote",
                    bullets=[
                        Bullet(text="Led cloud security architecture reviews for 20+ services."),
                        Bullet(text="Reduced critical vulnerabilities by 40% through automated scanning."),
                    ],
                    priority=1,
                ),
                Position(
                    company="Beta Inc",
                    title="Security Analyst",
                    start_date="2019-06",
                    end_date="2021-12",
                    is_current=False,
                    location="New York, NY",
                    bullets=[
                        Bullet(text="Performed penetration testing on web applications."),
                    ],
                    priority=2,
                ),
            ]
        ),
        skills=Skills(
            categories=[
                SkillCategory(
                    label="Cloud Security",
                    items=["AWS", "Azure", "GCP"],
                    priority=1,
                ),
                SkillCategory(
                    label="Languages",
                    items=["Python", "Go", "Bash"],
                    priority=2,
                ),
            ]
        ),
        education=Education(
            entries=[
                EducationEntry(
                    institution="MIT",
                    degree="B.Sc. Computer Science",
                    field="Cybersecurity",
                    start_date="2015-09",
                    end_date="2019-06",
                    gpa="3.9",
                    location="Cambridge, MA",
                    notes=["Dean's List 2017–2019"],
                    priority=1,
                )
            ]
        ),
        projects=Projects(
            projects=[
                Project(
                    name="CloudGuard",
                    description="Open-source cloud security posture management tool.",
                    technologies=["Python", "AWS SDK", "Terraform"],
                    url="https://github.com/janedoe/cloudguard",
                    bullets=[Bullet(text="Scans 50+ AWS services for misconfigurations.")],
                    priority=1,
                )
            ]
        ),
        certifications=Certifications(
            certifications=[
                Certification(
                    name="AWS Certified Security – Specialty",
                    issuer="Amazon Web Services",
                    date="2023-05",
                    expiry="2026-05",
                )
            ]
        ),
        meta=Meta(),
        locale=locale,
        template_name="classic",
        output_format="pdf",
    )


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------


class TestGetExporter:
    def test_md(self) -> None:
        exp = get_exporter("md")
        assert isinstance(exp, MarkdownExporter)

    def test_markdown_alias(self) -> None:
        exp = get_exporter("markdown")
        assert isinstance(exp, MarkdownExporter)

    def test_docx(self) -> None:
        exp = get_exporter("docx")
        assert isinstance(exp, DocxExporter)

    def test_pdf(self) -> None:
        exp = get_exporter("pdf")
        assert isinstance(exp, PdfExporter)

    def test_case_insensitive(self) -> None:
        assert isinstance(get_exporter("PDF"), PdfExporter)
        assert isinstance(get_exporter("DOCX"), DocxExporter)
        assert isinstance(get_exporter("MD"), MarkdownExporter)

    def test_dot_prefix_stripped(self) -> None:
        assert isinstance(get_exporter(".pdf"), PdfExporter)

    def test_unknown_raises(self) -> None:
        with pytest.raises(ValueError, match="Unknown export format"):
            get_exporter("xlsx")


# ---------------------------------------------------------------------------
# BaseExporter._resolve_output
# ---------------------------------------------------------------------------


class TestResolveOutput:
    def test_adds_extension(self, tmp_path: Path) -> None:
        exp = MarkdownExporter()
        result = exp._resolve_output(tmp_path / "resume")
        assert result.suffix == ".md"

    def test_preserves_correct_extension(self, tmp_path: Path) -> None:
        exp = MarkdownExporter()
        result = exp._resolve_output(tmp_path / "resume.md")
        assert result.suffix == ".md"

    def test_replaces_wrong_extension(self, tmp_path: Path) -> None:
        exp = MarkdownExporter()
        result = exp._resolve_output(tmp_path / "resume.pdf")
        assert result.suffix == ".md"

    def test_creates_parent_dirs(self, tmp_path: Path) -> None:
        exp = MarkdownExporter()
        deep = tmp_path / "a" / "b" / "c" / "resume"
        exp._resolve_output(deep)
        assert deep.parent.is_dir()


# ---------------------------------------------------------------------------
# MarkdownExporter
# ---------------------------------------------------------------------------


class TestMarkdownExporter:
    def test_export_creates_file(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        assert out.exists()
        assert out.suffix == ".md"

    def test_output_contains_name(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        content = out.read_text(encoding="utf-8")
        assert "Jane Doe" in content

    def test_output_contains_experience(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        content = out.read_text(encoding="utf-8")
        assert "Acme Corp" in content
        assert "Senior Security Engineer" in content

    def test_output_contains_skills(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        content = out.read_text(encoding="utf-8")
        assert "Cloud Security" in content
        assert "AWS" in content

    def test_output_contains_education(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        content = out.read_text(encoding="utf-8")
        assert "MIT" in content

    def test_output_contains_certifications(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        content = out.read_text(encoding="utf-8")
        assert "AWS Certified Security" in content

    def test_no_html_escaping(self, tmp_path: Path) -> None:
        """Markdown output must not escape ampersands or angle brackets."""
        ctx = _make_context()
        ctx.profile.title = "Security & Risk Engineer"
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        content = out.read_text(encoding="utf-8")
        assert "&amp;" not in content
        assert "Security & Risk Engineer" in content

    def test_missing_template_raises(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        with pytest.raises(ExportError, match="Template file not found"):
            exp.export(ctx, tmp_path / "nonexistent_template", tmp_path / "out")

    def test_returns_absolute_path(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = MarkdownExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        assert out.is_absolute()

    def test_format_name(self) -> None:
        assert MarkdownExporter().format_name == "md"


# ---------------------------------------------------------------------------
# DocxExporter
# ---------------------------------------------------------------------------


class TestDocxExporter:
    def test_export_creates_file(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = DocxExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        assert out.exists()
        assert out.suffix == ".docx"

    def test_output_is_valid_docx(self, tmp_path: Path) -> None:
        """A valid .docx is a ZIP file starting with the PK magic bytes."""
        ctx = _make_context()
        exp = DocxExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        with open(out, "rb") as fh:
            magic = fh.read(2)
        assert magic == b"PK", "DOCX file must be a valid ZIP/PK archive"

    def test_output_contains_name(self, tmp_path: Path) -> None:
        from docx import Document  # type: ignore[import-untyped]

        ctx = _make_context()
        exp = DocxExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text

    def test_output_contains_experience(self, tmp_path: Path) -> None:
        from docx import Document  # type: ignore[import-untyped]

        ctx = _make_context()
        exp = DocxExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in full_text

    def test_no_unicode_bullets(self, tmp_path: Path) -> None:
        """DOCX must not contain raw unicode bullet characters (•, ▪, ◦)."""
        ctx = _make_context()
        exp = DocxExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        raw = out.read_bytes().decode("utf-8", errors="replace")
        for char in ("•", "▪", "◦", "·"):
            assert char not in raw, f"Unicode bullet {char!r} found in DOCX"

    def test_returns_absolute_path(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = DocxExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        assert out.is_absolute()

    def test_format_name(self) -> None:
        assert DocxExporter().format_name == "docx"

    def test_empty_sections_dont_crash(self, tmp_path: Path) -> None:
        """Exporter must handle a context with all optional sections empty."""
        ctx = ResumeContext(
            profile=Profile(name="Min Imum"),
            experience=Experience(),
            skills=Skills(),
            education=Education(),
            projects=Projects(),
            certifications=Certifications(),
            meta=Meta(),
        )
        exp = DocxExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume_min")
        assert out.exists()


# ---------------------------------------------------------------------------
# PdfExporter
# ---------------------------------------------------------------------------


class TestPdfExporter:
    def test_export_creates_file(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = PdfExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        assert out.exists()
        assert out.suffix == ".pdf"

    def test_output_is_valid_pdf(self, tmp_path: Path) -> None:
        """A valid PDF starts with the %PDF magic bytes."""
        ctx = _make_context()
        exp = PdfExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        with open(out, "rb") as fh:
            magic = fh.read(4)
        assert magic == b"%PDF", "Output is not a valid PDF file"

    def test_output_is_non_empty(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = PdfExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        assert out.stat().st_size > 1024, "PDF file is suspiciously small"

    def test_returns_absolute_path(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = PdfExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume")
        assert out.is_absolute()

    def test_format_name(self) -> None:
        assert PdfExporter().format_name == "pdf"

    def test_rtl_locale(self, tmp_path: Path) -> None:
        """RTL locale (Hebrew) must not crash the PDF exporter."""
        ctx = _make_context(locale="he")
        exp = PdfExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume_he")
        assert out.exists()
        with open(out, "rb") as fh:
            assert fh.read(4) == b"%PDF"

    def test_missing_template_raises(self, tmp_path: Path) -> None:
        ctx = _make_context()
        exp = PdfExporter()
        with pytest.raises(ExportError, match="HTML template not found"):
            exp.export(ctx, tmp_path / "nonexistent_template", tmp_path / "out")

    def test_empty_sections_dont_crash(self, tmp_path: Path) -> None:
        ctx = ResumeContext(
            profile=Profile(name="Min Imum"),
            experience=Experience(),
            skills=Skills(),
            education=Education(),
            projects=Projects(),
            certifications=Certifications(),
            meta=Meta(),
        )
        exp = PdfExporter()
        out = exp.export(ctx, TEMPLATE_DIR, tmp_path / "resume_min")
        assert out.exists()
