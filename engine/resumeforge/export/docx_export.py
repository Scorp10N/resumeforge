"""DOCX exporter — builds a Word document programmatically via python-docx."""

from __future__ import annotations

import logging
from pathlib import Path

from resumeforge.data.schema import Position, ResumeContext
from resumeforge.export.base import BaseExporter, ExportError

logger = logging.getLogger(__name__)


class DocxExporter(BaseExporter):
    @property
    def format_name(self) -> str:
        return "docx"

    def export(
        self,
        context: ResumeContext,
        template_dir: Path,
        output_path: Path,
    ) -> Path:
        output_path = self._resolve_output(output_path)

        try:
            from docx import Document  # type: ignore[import-untyped]
            from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ExportError("python-docx not installed") from exc

        try:
            doc = Document()

            # --- Page margins ---
            for section in doc.sections:
                section.top_margin = Pt(36)
                section.bottom_margin = Pt(36)
                section.left_margin = Pt(54)
                section.right_margin = Pt(54)

            p = context.profile

            # --- Header: Name ---
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(p.name)
            name_run.bold = True
            name_run.font.size = Pt(20)

            # --- Header: Title ---
            if p.title:
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(p.title)
                title_run.font.size = Pt(11)

            # --- Header: Contact line ---
            contact_parts = [x for x in [p.email, p.phone, p.linkedin, p.location] if x]
            if contact_parts:
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_para.add_run(" | ".join(contact_parts)).font.size = Pt(9)

            # --- Summary ---
            if p.summary:
                _add_section_heading(doc, "Professional Summary")
                doc.add_paragraph(p.summary)

            # --- Experience ---
            if context.experience.positions:
                _add_section_heading(doc, "Experience")
                for pos in sorted(context.experience.positions, key=lambda x: x.priority):
                    _add_position(doc, pos)

            # --- Skills ---
            if context.skills.categories:
                _add_section_heading(doc, "Core Competencies")
                for cat in sorted(context.skills.categories, key=lambda x: x.priority):
                    p_skills = doc.add_paragraph()
                    run = p_skills.add_run(f"{cat.label}: ")
                    run.bold = True
                    p_skills.add_run(", ".join(cat.items))

            # --- Education ---
            if context.education.entries:
                _add_section_heading(doc, "Education")
                for edu in sorted(context.education.entries, key=lambda x: x.priority):
                    edu_para = doc.add_paragraph()
                    edu_para.add_run(f"{edu.degree} – {edu.institution}").bold = True
                    if edu.start_date:
                        edu_para.add_run(f"  {edu.start_date} – {edu.end_date or 'Present'}")

            # --- Projects ---
            if context.projects.projects:
                _add_section_heading(doc, "Projects")
                for proj in sorted(context.projects.projects, key=lambda x: x.priority):
                    proj_para = doc.add_paragraph()
                    proj_para.add_run(proj.name + ": ").bold = True
                    proj_para.add_run(proj.description)
                    for b in proj.bullets:
                        doc.add_paragraph(b.text, style="List Bullet")

            # --- Certifications ---
            if context.certifications.certifications:
                _add_section_heading(doc, "Certifications")
                for cert in context.certifications.certifications:
                    cert_para = doc.add_paragraph()
                    cert_para.add_run(cert.name).bold = True
                    meta_parts = [cert.issuer, cert.date]
                    if cert.expiry:
                        meta_parts.append(f"Expires {cert.expiry}")
                    cert_para.add_run("  —  " + " · ".join(p for p in meta_parts if p))

            doc.save(str(output_path))
            logger.info("DOCX export → %s", output_path)
            return output_path

        except ExportError:
            raise
        except Exception as exc:
            raise ExportError(f"DOCX export failed: {exc}") from exc


def _add_section_heading(doc: object, text: str) -> None:
    from docx.shared import Pt  # type: ignore[import-untyped]

    heading = doc.add_paragraph()  # type: ignore[union-attr]
    run = heading.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    # Add bottom border via paragraph style could be done with XML — keeping simple for v1
    doc.add_paragraph()  # type: ignore[union-attr]


def _add_position(doc: object, pos: Position) -> None:  # type: ignore[type-arg]
    from docx.shared import Pt  # type: ignore[import-untyped]

    header = doc.add_paragraph()  # type: ignore[union-attr]
    header.add_run(f"{pos.title}  –  {pos.company}").bold = True
    dates = f"{pos.start_date} – {'Present' if pos.is_current else (pos.end_date or '')}"
    header.add_run(f"  |  {dates}").font.size = Pt(9)

    for bullet in pos.bullets:
        doc.add_paragraph(bullet.text, style="List Bullet")  # type: ignore[union-attr]
